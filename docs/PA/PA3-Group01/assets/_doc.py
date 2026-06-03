# -*- coding: utf-8 -*-
"""Shared python-docx helpers for StoryLens PA3 documents."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INK = RGBColor(0x1a, 0x14, 0x0e)
ACCENT = RGBColor(0xb9, 0x1c, 0x1c)   # crimson
NAVY = RGBColor(0x16, 0x32, 0x5c)
TEAL = RGBColor(0x0f, 0x76, 0x6e)
GREY = RGBColor(0x55, 0x55, 0x55)
HEADER_FILL = "1f2a44"   # navy header for tables
SUBFILL = "e9eef6"
FONT = "Times New Roman"


def new_doc():
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(12)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    for name, size, color, bold in [
        ("Heading 1", 16, NAVY, True),
        ("Heading 2", 13.5, ACCENT, True),
        ("Heading 3", 12.5, INK, True),
    ]:
        s = doc.styles[name]
        s.font.name = FONT
        s.font.size = Pt(size)
        s.font.color.rgb = color
        s.font.bold = bold
    # page margins
    for sec in doc.sections:
        sec.top_margin = Cm(2.2); sec.bottom_margin = Cm(2.2)
        sec.left_margin = Cm(2.4); sec.right_margin = Cm(2.0)
    return doc


def _shade(cell, hexfill):
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), hexfill)
    cell._tc.get_or_add_tcPr().append(sh)


def _set_cell_font(cell, size=10.5, bold=False, color=None, white=False):
    for p in cell.paragraphs:
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        for r in p.runs:
            r.font.name = FONT
            r.font.size = Pt(size)
            r.font.bold = bold
            if white:
                r.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
            elif color is not None:
                r.font.color.rgb = color


def H1(doc, text):
    return doc.add_heading(text, level=1)


def H2(doc, text):
    return doc.add_heading(text, level=2)


def H3(doc, text):
    return doc.add_heading(text, level=3)


def P(doc, text, bold=False, italic=False, size=12, align=None, after=6, before=0, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.line_spacing = 1.18
    if align == "center": p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if align == "right": p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if align == "justify": p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.name = FONT; r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color is not None: r.font.color.rgb = color
    return p


def runs(doc, parts, align="justify", after=6, size=12):
    """parts: list of (text, {bold,italic,color}) tuples for mixed formatting."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.18
    if align == "justify": p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for text, fmt in parts:
        r = p.add_run(text)
        r.font.name = FONT; r.font.size = Pt(size)
        r.bold = fmt.get("bold", False); r.italic = fmt.get("italic", False)
        if "color" in fmt: r.font.color.rgb = fmt["color"]
    return p


def bullet(doc, text, level=0, bold_lead=None, size=12):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    if bold_lead:
        r = p.add_run(bold_lead); r.bold = True; r.font.name = FONT; r.font.size = Pt(size)
    r = p.add_run(text); r.font.name = FONT; r.font.size = Pt(size)
    return p


def numbered(doc, text, size=12):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.font.name = FONT; r.font.size = Pt(size)
    return p


def table(doc, headers, rows, widths=None, header_fill=HEADER_FILL, fontsize=10.5, zebra=True):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        _shade(hdr[i], header_fill)
        _set_cell_font(hdr[i], size=fontsize, bold=True, white=True)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            _set_cell_font(cells[i], size=fontsize)
            if zebra and ri % 2 == 1:
                _shade(cells[i], "f3f5f9")
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    # spacing after
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(8)
    p.paragraph_format.space_after = Pt(6); p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.line_spacing = 1.0
    pPr = p._p.get_or_add_pPr()
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), "f4f1ea")
    pPr.append(sh)
    r = p.add_run(text)
    r.font.name = "Consolas"; r.font.size = Pt(9.0)
    return p


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text); r.italic = True; r.font.name = FONT; r.font.size = Pt(10.5)
    r.font.color.rgb = GREY
    return p


def figure(doc, path, width=6.3, cap=None):
    doc.add_picture(path, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if cap: caption(doc, cap)


def placeholder(doc, route, what, height_lines=10):
    """A bordered box where the user pastes a screenshot."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    cell.width = Inches(6.2)
    _shade(cell, "f6f3ec")
    p0 = cell.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p0.add_run("[  CHÈN ẢNH CHỤP MÀN HÌNH TẠI ĐÂY  ]")
    r.bold = True; r.font.name = FONT; r.font.size = Pt(11); r.font.color.rgb = ACCENT
    for _ in range(max(1, height_lines)):
        cell.add_paragraph()
    p2 = cell.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"Route: {route}")
    r2.font.name = "Consolas"; r2.font.size = Pt(9.5); r2.font.color.rgb = NAVY
    p3 = cell.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(f"Cần chụp: {what}")
    r3.italic = True; r3.font.name = FONT; r3.font.size = Pt(9.5); r3.font.color.rgb = GREY
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def page_break(doc):
    doc.add_page_break()


def hrule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), "1f2a44")
    pbdr.append(bottom); pPr.append(pbdr)
